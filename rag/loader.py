# rag/loader.py
from __future__ import annotations
from pathlib import Path
from typing import List, Tuple
import logging

from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

# pypdf 对不规范 PDF 会刷大量无害警告(如 "Multiple definitions in dictionary"),
# 不影响文本抽取,这里压制到 ERROR 级别。
logging.getLogger("pypdf").setLevel(logging.ERROR)
logging.getLogger("pdfminer").setLevel(logging.ERROR)


def _table_to_markdown(table: list) -> str:
    """把 pdfplumber 抽出的二维表格转成 Markdown 表格,保留行列结构。"""
    rows = [[("" if c is None else str(c).replace("\n", " ").strip()) for c in row]
            for row in table if row]
    rows = [r for r in rows if any(cell for cell in r)]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * width) + " |"
    body = ["| " + " | ".join(r) + " |" for r in rows[1:]]
    return "\n".join([header, sep, *body])


def _detect_columns(page) -> int:
    """检测页面是否为双栏排版:统计跨中线的词占比,极少跨中线且左右两侧均有
    足量文字 → 判定为双栏。学术/诊疗指南类 PDF 常见双栏,跨栏逐行抽取会把
    中文句子打散成乱码,必须分栏。"""
    try:
        words = page.extract_words()
    except Exception:  # noqa: BLE001
        return 1
    if len(words) < 40:               # 文字太少(封面/图页),按单栏处理
        return 1
    w = page.width
    center = w / 2
    margin = w * 0.04
    crossing = sum(1 for d in words if d["x0"] < center - margin and d["x1"] > center + margin)
    left = sum(1 for d in words if (d["x0"] + d["x1"]) / 2 < center)
    right = len(words) - left
    if crossing / len(words) < 0.06 and left > len(words) * 0.25 and right > len(words) * 0.25:
        return 2
    return 1


def _extract_page_text(page) -> str:
    """按栏抽取页面正文:双栏则先左后右分别抽取再拼接,单栏直接抽。"""
    if _detect_columns(page) == 2:
        w, h = page.width, page.height
        parts = []
        for i in range(2):
            crop = page.crop((w * i / 2, 0, w * (i + 1) / 2, h))
            parts.append(crop.extract_text() or "")
        return "\n".join(parts)
    return page.extract_text() or ""


def _read_pdf(path: Path) -> str:
    """表格感知 + 栏感知 PDF 解析:自动识别双栏并分栏抽取(避免中文被跨栏打散),
    表格单独转为 Markdown 附在页尾。pdfplumber 失败时回退到 pypdf 纯文本。"""
    try:
        import pdfplumber
        parts: List[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                parts.append(_extract_page_text(page))
                for table in (page.extract_tables() or []):
                    md = _table_to_markdown(table)
                    if md:
                        parts.append("\n[表格]\n" + md)
        text = "\n".join(parts).strip()
        if text:
            return text
        logger.warning("pdfplumber 未抽到文本,回退 pypdf: %s", path.name)
    except Exception as e:  # noqa: BLE001
        logger.warning("pdfplumber 解析失败,回退 pypdf %s: %s", path.name, e)
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str:
    """抽段落正文 + 把 Word 表格转为 Markdown 表格,保留结构。"""
    doc = Document(str(path))
    parts: List[str] = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in tbl.rows]
        md = _table_to_markdown(rows)
        if md:
            parts.append("\n[表格]\n" + md)
    return "\n".join(parts)


def load_documents(docs_dir: Path) -> List[Tuple[str, str]]:
    """返回 [(全文, 文件名)],解析失败的文件跳过并记日志。"""
    results: List[Tuple[str, str]] = []
    docs_dir = Path(docs_dir)
    for path in sorted(docs_dir.glob("*")):
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                text = _read_pdf(path)
            elif ext == ".docx":
                text = _read_docx(path)
            elif ext in (".txt", ".md"):
                # 纯文本/Markdown(如扫描版 PDF 经 OCR 离线转出的正文)直接读入。
                text = path.read_text(encoding="utf-8", errors="ignore")
            else:
                continue
        except Exception as e:  # noqa: BLE001
            logger.warning("解析失败,跳过 %s: %s", path.name, e)
            continue
        if text.strip():
            results.append((text, path.name))
    return results
