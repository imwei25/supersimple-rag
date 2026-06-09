# tests/test_loader.py
from pathlib import Path
from docx import Document
from rag.loader import load_documents

def test_load_docx(tmp_path):
    p = tmp_path / "a.docx"
    doc = Document()
    doc.add_paragraph("这是第一段中文。")
    doc.add_paragraph("这是第二段。")
    doc.save(p)
    results = load_documents(tmp_path)
    assert len(results) == 1
    text, source = results[0]
    assert "第一段" in text
    assert source == "a.docx"

def test_skips_unknown_extension(tmp_path):
    (tmp_path / "note.xyz").write_text("ignored", encoding="utf-8")
    (tmp_path / "pic.png").write_bytes(b"\x89PNG")
    assert load_documents(tmp_path) == []

def test_loads_txt_and_md(tmp_path):
    # .txt/.md 受支持:扫描版 PDF 经 OCR 离线转出的正文据此入库
    (tmp_path / "ocr.txt").write_text("正文内容", encoding="utf-8")
    (tmp_path / "doc.md").write_text("# 标题\n正文", encoding="utf-8")
    loaded = dict((src, txt) for txt, src in load_documents(tmp_path))
    assert loaded["ocr.txt"] == "正文内容"
    assert "标题" in loaded["doc.md"]
